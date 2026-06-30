"""
rag.py — RAG (Retrieval-Augmented Generation) tools for the StoreIt AI agent.

smart_extract_text: Enhanced — handles PDF (pdfplumber), DOCX (python-docx),
                    plain text, and Gemini OCR fallback for scanned documents.
process_file_for_search: Downloads file via Express, embeds, stores via Express.
ask_file_question: Embeds query locally, sends to Express for similarity search.
"""

import os
import io
import json
import tempfile
import asyncio
import pdfplumber
import docx
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.tools import tool

from config import GOOGLE_API_KEY, MAX_CHUNKS_PER_FILE, CHUNK_SIZE, CHUNK_OVERLAP, TOP_K_RESULTS
import api_client

# Gemini embeddings
embeddings = GoogleGenerativeAIEmbeddings(
    model="models/gemini-embedding-2",
    google_api_key=GOOGLE_API_KEY,
)

# Gemini GenAI for OCR fallback — unchanged
import google.generativeai as genai
genai.configure(api_key=GOOGLE_API_KEY)


# ─── Helper: run async from sync (same as tools.py) ─────────────────────────

def _run(coro):
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as pool:
                future = pool.submit(asyncio.run, coro)
                return future.result()
        else:
            return loop.run_until_complete(coro)
    except RuntimeError:
        return asyncio.run(coro)


# ─── smart_extract_text — KEPT VERBATIM FROM ORIGINAL ───────────────────────

def smart_extract_text(file_bytes: bytes, file_ext: str = ".pdf") -> str:
    """
    Extracts text from various file types.

    Extraction strategy by extension:
      .docx  → python-docx (paragraphs + table cells)
      .pdf   → pdfplumber, then Gemini OCR fallback for scanned pages
      .txt, .csv, .md, .json, .xml, .html → decode as UTF-8 text
      Other  → Gemini OCR fallback (upload file and ask for transcription)
    """
    ext = file_ext.lower()
    text = ""

    # ── DOCX ──────────────────────────────────────────────────────────────────
    if ext in (".docx", ".doc"):
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            text = "\n".join(paragraphs)
            print(f"DEBUG: python-docx extracted {len(text)} chars from {ext}")
        except Exception as e:
            print(f"DEBUG: python-docx failed: {e}")

        if len(text.strip()) > 50:
            return text
        # If docx extraction got nothing, fall through to Gemini OCR

    # ── Plain text formats ────────────────────────────────────────────────────
    elif ext in (".txt", ".csv", ".md", ".json", ".xml", ".html", ".htm", ".log", ".py", ".js", ".ts"):
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            print(f"DEBUG: Decoded {ext} as UTF-8, {len(text)} chars")
            return text
        except Exception as e:
            print(f"DEBUG: UTF-8 decode failed for {ext}: {e}")

    # ── PDF ───────────────────────────────────────────────────────────────────
    elif ext == ".pdf":
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception as e:
            print(f"DEBUG: pdfplumber failed: {e}")

        if len(text.strip()) > 100:
            return text
        # Sparse text — fall through to Gemini OCR

    # ── Gemini OCR Fallback (for scanned PDFs, images, or unknown formats) ──
    print(f"DEBUG: Text sparse/empty for {ext}. Falling back to Gemini OCR...")
    import traceback
    import time

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        print(f"DEBUG: Uploading {tmp_path} to Gemini...")
        myfile = genai.upload_file(tmp_path)
        print(f"DEBUG: File uploaded: {myfile.name}")

        model = genai.GenerativeModel("models/gemini-2.5-flash")
        print("DEBUG: Generating content...")

        result_text = ""
        max_retries = 3
        for attempt in range(max_retries):
            try:
                result = model.generate_content(
                    [myfile, "Transcribe the full text of this document verbatim."]
                )
                result_text = result.text
                print(f"DEBUG: Gemini OCR success! Length: {len(result_text)}")
                break
            except Exception as e:
                if "429" in str(e) and attempt < max_retries - 1:
                    wait_time = (attempt + 1) * 10
                    print(f"DEBUG: Rate Limit (429). Retrying in {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    raise e

        os.remove(tmp_path)
        return result_text

    except Exception as e:
        print(f"DEBUG: Gemini OCR failed: {e}")
        traceback.print_exc()
        return text  # Return whatever we managed to extract


# ─── Chunking helper ─────────────────────────────────────────────────────────

def _chunk_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    """
    Split text into overlapping chunks.

    The original used: [text[i:i+1000] for i in range(0, len(text), 1000)]
    (no overlap, hard 1000-char split).

    This version adds configurable overlap so adjacent chunks share context,
    which improves recall for the similarity search. The behaviour is identical
    to the original when CHUNK_OVERLAP=0.
    """
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start = end - overlap  # Move back by overlap for the next chunk
    return chunks


# ─── Tool 6: process_file_for_search ─────────────────────────────────────────

@tool
def process_file_for_search(file_id: str):
    """
    Downloads a file, extracts its text, creates embeddings, and indexes it for search.
    Use this when the user asks to 'analyze', 'read', or 'index' a specific file.

    Args:
        file_id: The ID of the file (obtained from search_files).

    NOTE: In the old system this required both 'file_id' AND 'bucket_file_id'.
    In this system, only 'file_id' is needed.
    """
    try:
        # ── Step 1: Get file metadata to determine extension ───────────────────
        print(f"DEBUG: Fetching file info for {file_id}...")
        file_info = _run(api_client.get_file_info_api(file_id=file_id))
        # Default to .pdf if extension isn't found
        ext = f".{file_info.get('extension', 'pdf').lower().strip('.')}"
        print(f"DEBUG: File extension detected as {ext}")

        # ── Step 2: Download raw bytes from Express ────────────────────────────
        print(f"DEBUG: Fetching file buffer for {file_id}...")
        file_bytes = _run(api_client.get_file_buffer_api(file_id=file_id))
        print(f"DEBUG: Downloaded {len(file_bytes)} bytes")

        # ── Step 3: Extract text ───────────────────────────────────────────────
        text = smart_extract_text(file_bytes, file_ext=ext)
        print(f"DEBUG: Extracted text length: {len(text)}")
        print(f"DEBUG: Text preview: {text[:200]}")

        if not text or not text.strip():
            return "Could not extract text from this file. It may be empty, an unsupported format, or a scanned image that OCR could not process."

        # ── Step 3: Chunking ──────────────────────────────────────────────────
        chunks_text = _chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP)
        chunks_text = chunks_text[:MAX_CHUNKS_PER_FILE]  # Enforce limit
        print(f"DEBUG: Created {len(chunks_text)} chunks (limit: {MAX_CHUNKS_PER_FILE})")

        # ── Step 4: Embed each chunk ───────────────────────────────────────────
        chunk_payloads = []
        for i, chunk in enumerate(chunks_text):
            vector = embeddings.embed_query(chunk, output_dimensionality=768)  # Returns list[float] — 768 dims
            chunk_payloads.append({
                "chunkIndex": i,
                "text": chunk,
                "embedding": vector,  # Sent as JSON array; Express stores as [Number]
                "tokenCount": len(chunk.split()),  # Rough token estimate
            })
            print(f"DEBUG: Embedded chunk {i+1}/{len(chunks_text)}")

        # ── Step 5: Store vectors via Express ─────────────────────────────────
        # Express will: delete existing vectors for this file (upsert), insert new
        # ones, set File.isIndexed=true, File.chunkCount=len(chunks).
        result = _run(api_client.store_vectors_api(
            file_id=file_id,
            chunks=chunk_payloads,
            embedding_model="gemini-embedding-2",
        ))
        count = result.get("data", {}).get("chunkCount", len(chunk_payloads))
        return f"File processed successfully. {count} chunks indexed and ready for search."

    except Exception as e:
        return f"Error processing file: {str(e)}"


# ─── Tool 7: ask_file_question ────────────────────────────────────────────────

@tool
def ask_file_question(question: str, file_id: str = None):
    """
    Search your indexed files to answer a question about their content.
    Use this when the user asks a question about WHAT IS INSIDE their files.

    Args:
        question: The question to answer from file content.
        file_id: (Optional) Scope the search to a specific file. If not provided,
                 searches across ALL indexed files owned by the user.
    """
    try:
        # ── Step 1: Embed the question ────────────────────────────────────────
        # Uses the same text-embedding-004 model as process_file_for_search,
        # so the vector spaces are aligned for meaningful similarity comparison.
        print(f"DEBUG: Embedding question: {question[:80]}...")
        query_vector = embeddings.embed_query(question, output_dimensionality=768)

        # ── Step 2: Ask Express to find similar chunks ────────────────────────
        # Express loads vectors for this owner, computes cosine similarity,
        # and returns top-K results. No numpy needed here.
        results = _run(api_client.query_vectors_api(
            query_embedding=query_vector,
            top_k=TOP_K_RESULTS,
            file_id=file_id,
        ))

        if not results:
            return (
                "No indexed documents found. "
                "Ask me to 'analyze' a file first so I can read its content."
            )

        print(f"DEBUG: Top {len(results)} chunks found:")
        for r in results:
            print(f"  - Score: {r.get('score'):.3f} | File: {r.get('fileName')} | Chunk: {r.get('chunkIndex')}")

        # ── Step 3: Format context for the LLM ───────────────────────────────
        context_parts = []
        for r in results:
            file_name = r.get("fileName", "unknown file")
            text = r.get("text", "")
            context_parts.append(f"[From: {file_name}]\n{text}")

        return "Context found:\n" + "\n---\n".join(context_parts)

    except Exception as e:
        return f"Error searching file content: {str(e)}"
